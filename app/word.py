"""The reading record as a Word file.

The record's structure lives in one place, `app/templates/export.md`, and it stays there: this
module takes that rendered markdown and walks it into a document. Duplicating the structure in a
second builder is how the two versions of a record start disagreeing about what the reading found.

python-docx is already a dependency (intake reads .docx with it), and the styles used here are
Word's own — Title, Heading 1–3, Normal, List Bullet, Quote — so the file opens looking like a
Word document and not like someone's idea of one.
"""
from __future__ import annotations

import io
import re

# `[text](#somewhere)` — a link in a document that has no links is its text.
_LINK = re.compile(r"\[([^\]]*)\]\([^)]*\)")


def _write(p, text: str) -> None:
    """One markdown line into one paragraph. `**bold**` is the only inline mark the record uses;
    `[S014]` is a passage id and stays exactly as it is."""
    for i, part in enumerate(_LINK.sub(r"\1", text).split("**")):
        if part:
            p.add_run(part).bold = bool(i % 2)


def document(markdown: str, title: str, when: str) -> bytes:
    """The rendered record, as .docx bytes."""
    import docx

    doc = docx.Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(when)
    doc.add_paragraph("Aperture reading record")
    doc.add_page_break()

    for line in markdown.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("#"):
            hashes = len(s) - len(s.lstrip("#"))
            # Five levels, because the record has five: the two groups of themes, a theme's name
            # and "materials where this theme appears" flattened onto one level make Word's
            # navigation pane useless.
            _write(doc.add_heading("", min(hashes, 5)), s[hashes:].strip())
        elif s.startswith("- "):
            _write(doc.add_paragraph(style="List Bullet"), s[2:])
        elif s.startswith("> "):
            _write(doc.add_paragraph(style="Quote"), s[2:])
        else:
            _write(doc.add_paragraph(), s)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
